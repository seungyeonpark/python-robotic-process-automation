import docx

doc = docx.Document()

par0 = doc.add_paragraph('add paragraph0')
par1 = par0.insert_paragraph_before('add paragraph1')

table = doc.add_table(rows=1, cols=3, style="Table Grid")
table.cell(0,0).text = "(0,0)"
table.rows[0].cells[1].text = "(0,1)"
table.columns[2].cells[0].text = "(0,2)"

for i in range(10):
    doc.add_heading('Level {}'.format(i), level=i)

doc.save('./output/docx_write.docx')
