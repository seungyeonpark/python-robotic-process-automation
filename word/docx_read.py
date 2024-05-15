import docx

doc = docx.Document('./output/docx_write.docx')

doc.paragraphs[0].text = 'Change first paragraph'

table = doc.tables[0]
new_row = table.add_row()
r_num = len(table.rows)
c_num = len(table.columns)

for c in range(c_num):
    new_row.cells[c].text = "Add ({}, {}) cell".format(r_num, c)

for p in doc.paragraphs:
    print(p.text)

for r in range(r_num):
    for c in range(c_num):
        print(table.cell(r, c).text)

doc.save('./output/docx_read.docx')
